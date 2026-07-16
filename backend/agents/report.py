import os
import logging
import docx
from typing import List, Dict, Any
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from backend.data.models import PaperMeta, GapClaim, Summary
from backend.clients.claude_client import ClaudeClient

logger = logging.getLogger("researchmind.report")

def generate_thematic_synthesis(query: str, summaries: List[Summary], gap_claims: List[GapClaim]) -> str:
    """
    Leverages LLM to synthesize a continuous, publication-grade academic survey of the literature.
    """
    logger.info("Generating thematic academic synthesis of the literature...")
    claude = ClaudeClient()
    
    papers_input = ""
    for idx, s in enumerate(summaries):
        papers_input += f"Paper {idx+1}: {s.title}\nSummary & Methodology: {s.summary_text}\n\n"
        
    gaps_input = ""
    for idx, gap in enumerate(gap_claims):
        gaps_input += f"Gap {idx+1}: {gap.topic_label}\nDescription: {gap.description}\nDirections: {', '.join(gap.suggested_directions)}\n\n"
        
    prompt = f"""
Write a professional, publication-grade academic literature review and thematic synthesis on the topic: '{query}'.

Use the following analyzed papers as input data:
{papers_input}

Also incorporate discussion around these identified literature gaps:
{gaps_input}

Guidelines:
1. Write in a formal, objective, academic style (third-person).
2. Organize the content into at least three logical thematic subsections (e.g. 3.1 Methodological Paradigms, 3.2 Evaluation & Benchmarks, 3.3 Identified Limitations and Research Gaps).
3. Do NOT list the papers one-by-one. Instead, group, compare, and contrast their methods.
4. Cite papers using their authors or titles in standard format (e.g. "[FirstAuthor et al.]").
5. Connect the survey to the identified literature gaps, showing why current research has left these gaps open.
6. The synthesis should be comprehensive and flow naturally like a real journal survey paper (approx. 500-800 words).

Return ONLY the synthesized literature review text with standard Markdown headers. Do not include markdown code block backticks (like ```markdown).
"""
    try:
        response = claude.complete(
            prompt=prompt,
            system="You are an expert academic research writer and literature review synthesiser.",
            max_tokens=2500,
            temperature=0.2
        )
        return response.strip()
    except Exception as e:
        logger.error(f"Failed to generate thematic synthesis: {e}")
        # Fallback to simple concatenation
        fallback = "### 3.1 Literature Survey Overview\n"
        for s in summaries:
            fallback += f"**{s.title}**: {s.summary_text}\n\n"
        return fallback

def add_markdown_paragraphs_docx(doc, text: str):
    """
    Parses a simple markdown text block and adds it to python-docx Document.
    """
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=0)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)

def add_markdown_paragraphs_pdf(story, text: str, h1_style, h2_style, body_style):
    """
    Parses a simple markdown text block and adds it to ReportLab story.
    """
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("### "):
            story.append(Paragraph(line[4:], h2_style))
            story.append(Spacer(1, 4))
        elif line.startswith("## "):
            story.append(Paragraph(line[3:], h1_style))
            story.append(Spacer(1, 6))
        elif line.startswith("- "):
            story.append(Paragraph(f"• {line[2:]}", body_style))
            story.append(Spacer(1, 4))
        else:
            story.append(Paragraph(line, body_style))
            story.append(Spacer(1, 6))

def compile_markdown_draft(query: str, summaries: List[Summary], comparison_table: List[Dict[str, Any]], gap_claims: List[GapClaim], synthesis_text: str) -> str:
    """
    Compiles a plain text/markdown draft of the final report.
    """
    introduction = (
        f"# ResearchMind Literature Review Report\n\n"
        f"**Research Topic:** {query}\n\n"
        f"This report presents an automated synthesis of the current state of academic literature "
        f"on the topic. We analyzed {len(summaries)} relevant papers from arXiv and Semantic Scholar, "
        f"extracted structured methodology records, built a citation network, and detected candidate research gaps.\n"
    )
    
    table_section = "\n## 2. Comparison Matrix\n\n"
    table_section += "| Title | Year | Method | Dataset | Key Metric | Limitation | Status |\n"
    table_section += "|---|---|---|---|---|---|---|\n"
    for row in comparison_table:
        title_trunc = row['title'][:40] + "..." if len(row['title']) > 40 else row['title']
        table_section += f"| {title_trunc} | {row['year']} | {row['method']} | {row['dataset']} | {row['key_metric']} | {row['limitation']} | {row['verification_status']} |\n"
        
    summary_section = f"\n## 3. Thematic Literature Survey & Synthesis\n\n{synthesis_text}\n"
        
    gap_section = "\n## 4. Identified Research Gaps\n\n"
    if not gap_claims:
        gap_section += "No high-confidence research gaps were detected in this literature set.\n"
    else:
        for idx, gap in enumerate(gap_claims):
            gap_section += f"### Gap {idx+1}: {gap.topic_label}\n"
            gap_section += f"**Description:** {gap.description}\n\n"
            gap_section += f"**Citation Density:** {gap.citation_density:.2f} citations/paper\n\n"
            gap_section += "**Suggested Directions for Future Research:**\n"
            for dir_stmt in gap.suggested_directions:
                gap_section += f"- {dir_stmt}\n"
            gap_section += "\n"
            
    return introduction + table_section + summary_section + gap_section

def generate_docx(output_path: str, query: str, summaries: List[Summary], comparison_table: List[Dict[str, Any]], gap_claims: List[GapClaim], synthesis_text: str):
    """
    Generates a structured DOCX report.
    """
    doc = docx.Document()
    
    # Title
    doc.add_heading("ResearchMind Literature Review & Gap Analysis", 0)
    doc.add_paragraph(f"Focus Topic: {query}").bold = True
    
    # Intro
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        f"This report presents an automated literature review and research gap discovery "
        f"for the topic '{query}'. A total of {len(summaries)} papers were compiled and synthesized "
        f"from arXiv and Semantic Scholar databases."
    )
    
    # Comparison Table
    doc.add_heading("2. Comparison Matrix", level=1)
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Title'
    hdr_cells[1].text = 'Year'
    hdr_cells[2].text = 'Method'
    hdr_cells[3].text = 'Dataset'
    hdr_cells[4].text = 'Key Metric'
    hdr_cells[5].text = 'Limitation'
    hdr_cells[6].text = 'Status'
    
    for row in comparison_table:
        row_cells = table.add_row().cells
        row_cells[0].text = row['title'][:30] + "..." if len(row['title']) > 30 else row['title']
        row_cells[1].text = str(row['year'])
        row_cells[2].text = row['method']
        row_cells[3].text = row['dataset']
        row_cells[4].text = row['key_metric']
        row_cells[5].text = row['limitation']
        row_cells[6].text = row['verification_status']
        
    # Thematic Synthesis
    doc.add_heading("3. Thematic Literature Survey & Synthesis", level=1)
    add_markdown_paragraphs_docx(doc, synthesis_text)
        
    # Gaps
    doc.add_heading("4. Identified Research Gaps", level=1)
    if not gap_claims:
        doc.add_paragraph("No significant research gaps were identified in the corpus.")
    else:
        for idx, gap in enumerate(gap_claims):
            doc.add_heading(f"Gap {idx+1}: {gap.topic_label}", level=2)
            doc.add_paragraph(gap.description)
            doc.add_paragraph(f"Citation Density: {gap.citation_density:.2f} citations/paper")
            
            doc.add_paragraph("Suggested Future Directions:").bold = True
            for dir_stmt in gap.suggested_directions:
                doc.add_paragraph(dir_stmt, style='List Bullet')
                
    doc.save(output_path)
    logger.info(f"DOCX report saved to {output_path}")

def generate_pdf(output_path: str, query: str, summaries: List[Summary], comparison_table: List[Dict[str, Any]], gap_claims: List[GapClaim], synthesis_text: str):
    """
    Generates a premium PDF report using ReportLab.
    """
    doc = SimpleDocTemplate(output_path, pagesize=letter)
    styles = getSampleStyleSheet()
    
    # Custom Paragraph Styles
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Title'],
        fontName='Helvetica-Bold',
        fontSize=24,
        leading=28,
        textColor=colors.HexColor('#1E293B'),
        spaceAfter=15
    )
    
    h1_style = ParagraphStyle(
        'Heading1Style',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=16,
        leading=20,
        textColor=colors.HexColor('#0F172A'),
        spaceBefore=14,
        spaceAfter=8,
        keepWithNext=True
    )
    
    h2_style = ParagraphStyle(
        'Heading2Style',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=12,
        leading=16,
        textColor=colors.HexColor('#1E293B'),
        spaceBefore=10,
        spaceAfter=4,
        keepWithNext=True
    )
    
    body_style = ParagraphStyle(
        'BodyStyle',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        textColor=colors.HexColor('#334155'),
        spaceAfter=8
    )
    
    table_hdr_style = ParagraphStyle(
        'TableHdrStyle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=8,
        leading=10,
        textColor=colors.white
    )
    
    table_cell_style = ParagraphStyle(
        'TableCellStyle',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=8,
        leading=10,
        textColor=colors.HexColor('#334155')
    )
    
    story = []
    
    # Title
    story.append(Paragraph("Literature Review & Gap Analysis Report", title_style))
    story.append(Paragraph(f"<b>Research Focus:</b> {query}", body_style))
    story.append(Spacer(1, 10))
    
    # Section 1: Intro
    story.append(Paragraph("1. Introduction", h1_style))
    intro_p = (
        f"This report presents an automated systematic literature survey and candidate literature gap discovery "
        f"for the query '{query}'. A corpus of {len(summaries)} relevant papers was fetched, parsed, and "
        f"synthesized to analyze the state of current methodology implementations."
    )
    story.append(Paragraph(intro_p, body_style))
    story.append(Spacer(1, 10))
    
    # Section 2: Comparison Table
    story.append(Paragraph("2. Comparison Matrix", h1_style))
    
    # ReportLab table headers
    data = [[
        Paragraph("Title", table_hdr_style),
        Paragraph("Year", table_hdr_style),
        Paragraph("Method", table_hdr_style),
        Paragraph("Dataset", table_hdr_style),
        Paragraph("Metric", table_hdr_style),
        Paragraph("Limitation", table_hdr_style),
        Paragraph("Status", table_hdr_style)
    ]]
    
    # Table rows
    for row in comparison_table:
        title_short = row['title'][:25] + "..." if len(row['title']) > 25 else row['title']
        data.append([
            Paragraph(title_short, table_cell_style),
            Paragraph(str(row['year']), table_cell_style),
            Paragraph(row['method'], table_cell_style),
            Paragraph(row['dataset'], table_cell_style),
            Paragraph(row['key_metric'], table_cell_style),
            Paragraph(row['limitation'], table_cell_style),
            Paragraph(row['verification_status'], table_cell_style)
        ])
        
    t = Table(data, colWidths=[120, 35, 80, 80, 80, 80, 50])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1E293B')),
        ('ALIGN', (0,0), (-1,-1), 'LEFT'),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
        ('TOPPADDING', (0,0), (-1,-1), 6),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#E2E8F0')),
        ('BACKGROUND', (0,1), (-1,-1), colors.HexColor('#F8FAFC')),
    ]))
    
    story.append(t)
    story.append(Spacer(1, 15))
    story.append(PageBreak())
    
    # Section 3: Synthesis Survey
    story.append(Paragraph("3. Thematic Literature Survey & Synthesis", h1_style))
    add_markdown_paragraphs_pdf(story, synthesis_text, h1_style, h2_style, body_style)
    story.append(Spacer(1, 10))
    story.append(PageBreak())
    
    # Section 4: Gaps
    story.append(Paragraph("4. Identified Research Gaps", h1_style))
    if not gap_claims:
        story.append(Paragraph("No significant research gaps were identified in the corpus.", body_style))
    else:
        for idx, gap in enumerate(gap_claims):
            story.append(Paragraph(f"Gap {idx+1}: {gap.topic_label}", h2_style))
            story.append(Paragraph(gap.description, body_style))
            story.append(Paragraph(f"<b>Citation Density:</b> {gap.citation_density:.2f} citations/paper", body_style))
            story.append(Paragraph("<b>Suggested Future Directions:</b>", body_style))
            for dir_stmt in gap.suggested_directions:
                story.append(Paragraph(f"• {dir_stmt}", body_style))
            story.append(Spacer(1, 10))
            
    doc.build(story)
    logger.info(f"PDF report saved to {output_path}")

def run_report(state: dict) -> dict:
    """
    Compiles summaries and gap claims into PDF and DOCX documents.
    """
    query = state.get("query", "")
    summaries: List[Summary] = state.get("summaries", [])
    comparison_table: List[Dict[str, Any]] = state.get("comparison_table", [])
    gap_claims: List[GapClaim] = state.get("gap_claims", [])
    
    if "agent_status" not in state:
        state["agent_status"] = {}
        
    state["agent_status"]["report"] = "running"
    logger.info("Report Agent: Commencing report compilation.")
    
    # Call thematic synthesis generator
    synthesis_text = generate_thematic_synthesis(query, summaries, gap_claims)
    
    # Compile markdown draft for state reference
    state["report_draft"] = {
        "text": compile_markdown_draft(query, summaries, comparison_table, gap_claims, synthesis_text)
    }
    
    # Create base output directories
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    exports_dir = os.path.join(base_dir, "db", "exports")
    os.makedirs(exports_dir, exist_ok=True)
    
    # Output file paths
    pdf_path = os.path.join(exports_dir, "report.pdf")
    docx_path = os.path.join(exports_dir, "report.docx")
    
    try:
        generate_docx(docx_path, query, summaries, comparison_table, gap_claims, synthesis_text)
        state["report_draft"]["docx_path"] = docx_path
    except Exception as e:
        logger.error(f"Failed to generate DOCX report: {e}")
        
    try:
        generate_pdf(pdf_path, query, summaries, comparison_table, gap_claims, synthesis_text)
        state["report_draft"]["pdf_path"] = pdf_path
    except Exception as e:
        logger.error(f"Failed to generate PDF report: {e}")
        
    state["agent_status"]["report"] = "done"
    return state
